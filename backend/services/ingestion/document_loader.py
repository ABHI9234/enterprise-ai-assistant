import os
from pathlib import Path
from loguru import logger
from pypdf import PdfReader
from docx import Document
import markdown


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


def load_document(file_path: str) -> dict:
    """
    Load a document from disk and extract its text content.
    Returns a dict with filename, content, and file type.
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    extension = path.suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type: {extension}. "
            f"Supported types: {SUPPORTED_EXTENSIONS}"
        )

    logger.info(f"Loading document: {path.name} (type: {extension})")

    if extension == ".pdf":
        content = _load_pdf(path)
    elif extension == ".docx":
        content = _load_docx(path)
    elif extension == ".txt":
        content = _load_txt(path)
    elif extension == ".md":
        content = _load_markdown(path)

    logger.success(f"Successfully loaded: {path.name} ({len(content)} characters)")

    return {
        "filename": path.name,
        "file_path": str(path),
        "file_type": extension,
        "content": content,
        "char_count": len(content),
    }


def _load_pdf(path: Path) -> str:
    """Extract text from all pages of a PDF file."""
    reader = PdfReader(str(path))
    pages = []

    for page_num, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            pages.append(f"[Page {page_num + 1}]\n{text.strip()}")

    if not pages:
        raise ValueError(f"No extractable text found in PDF: {path.name}")

    return "\n\n".join(pages)


def _load_docx(path: Path) -> str:
    """Extract text from all paragraphs of a DOCX file."""
    doc = Document(str(path))
    paragraphs = []

    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                paragraphs.append(row_text)

    if not paragraphs:
        raise ValueError(f"No extractable text found in DOCX: {path.name}")

    return "\n\n".join(paragraphs)


def _load_txt(path: Path) -> str:
    """Read plain text file."""
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read().strip()

    if not content:
        raise ValueError(f"Text file is empty: {path.name}")

    return content


def _load_markdown(path: Path) -> str:
    """Convert markdown to plain text."""
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        md_content = f.read().strip()

    if not md_content:
        raise ValueError(f"Markdown file is empty: {path.name}")

    # Convert markdown to HTML then strip tags for clean text
    html = markdown.markdown(md_content)
    # Remove HTML tags manually for clean plain text
    import re
    clean = re.sub(r"<[^>]+>", " ", html)
    clean = re.sub(r"\s+", " ", clean).strip()

    return clean


def save_uploaded_file(file_bytes: bytes, filename: str, upload_dir: str) -> str:
    """
    Save uploaded file bytes to disk and return the saved file path.
    """
    upload_path = Path(upload_dir)
    upload_path.mkdir(parents=True, exist_ok=True)

    file_path = upload_path / filename

    with open(file_path, "wb") as f:
        f.write(file_bytes)

    logger.info(f"Saved uploaded file: {file_path}")
    return str(file_path)
